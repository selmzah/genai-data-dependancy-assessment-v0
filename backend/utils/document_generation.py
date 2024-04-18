from docx import Document
from docx.shared import Pt
import os


def create_documentation(general_explanation: str, json_data: list) -> str:
    """
    Permet la création d'un docx pour la rétrodocumentation
    :param json_data: Dictionnaire des explications des fonctions
    :type json_data: dict
    :param general_explanation: Explication générale de l'application
    :type general_explanation: str
    :return: Chemin du docx
    :rtype: str
    """
    # Création du dossier 'tmp' s'il n'existe pas
    tmp_dir = 'tmp'
    if not os.path.exists(tmp_dir):
        os.makedirs(tmp_dir)

    path_to_docx = os.path.join(tmp_dir, 'document.docx')

    document = Document()
    desired_font = 'Ubuntu'

    title = document.add_heading(level=1)
    run = title.add_run("Overall description")
    run.font.size = Pt(16)
    run.font.name = desired_font
    document.add_paragraph(general_explanation)

    for data in json_data:
        try :
            # Ajouter un titre général
            title = document.add_heading(level=1)
            run = title.add_run("General")
            run.font.size = Pt(16)
            run.font.name = desired_font

            # Définition des labels personnalisés
            custom_labels = {
                "file_name": "Nom du fichier",
                "langage": "Langage",
                "simple_explanation": "Objectif du script"
            }

            # Ajouter des informations générales
            general_info = data["general"]
            for info_key, info_value in general_info.items():
                label = custom_labels.get(info_key, info_key)
                para = document.add_paragraph()
                run = para.add_run(f"{label}: ")
                run.font.name = desired_font
                run.bold = True  # Mettre en gras le label
                run = para.add_run(info_value)
                run.font.name = desired_font  # Changer la typographie

            # Ajouter un titre d'explication
            title = document.add_heading(level=1)
            run = title.add_run("Explanation")
            run.font.size = Pt(16)
            run.font.name = desired_font  # Changer la typographie ici

            # Ajouter du texte d'explication
            explanation = document.add_paragraph(data["explanation_text"])
            for run in explanation.runs:
                run.font.name = desired_font  # Changer la typographie pour chaque 'run' dans le paragraphe

            # Ajouter un titre de tableau de fonction
            title = document.add_heading(level=1)
            run = title.add_run("Function Table")
            run.font.size = Pt(16)
            run.font.name = desired_font
            run.bold = True

            # Créer un tableau avec des bordures
            table = document.add_table(rows=1, cols=4, style='Table Grid')  # 'Table Grid' ajoute les bordures

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Function Name'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = 'Arguments'
            hdr_cells[3].text = 'Dependencies'

            # Mettre en gras les en-têtes du tableau
            for hdr_cell in hdr_cells:
                hdr_cell.paragraphs[0].runs[0].bold = True

            for fonction in data["fonction_tab"]:
                    row_cells = table.add_row().cells
                    row_cells[0].text = fonction["fonction_name"]
                    row_cells[1].text = fonction["description"]

                    # Ajoutez chaque argument comme un nouveau paragraphe dans la cellule
                    args_text = ''
                    for arg in fonction["arguments"]:
                        arg_text = f"{arg['arg_name']} ({arg['arg_type']}) : {arg['arg_description']}"
                        args_text += arg_text + "\n"

                    row_cells[2].text = args_text.strip()
                    row_cells[3].text = ', '.join(fonction["dependance"])
        except:
            print(fonction)

    # Enregistrez votre document
    document.save(path_to_docx)
    return path_to_docx
